# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SRC = r"D:/VsProject/CAC_DB/thesis_revision_source.docx"
OUT = r"D:/VsProject/CAC_DB/thesis_academic_revised.docx"


replacements = {
    36: "高等教育规模扩张与产业结构调整叠加，使高校毕业生就业压力与企业高质量岗位空缺并存，人才供需错配问题持续凸显。高校招生、教务、培养和就业等业务系统长期分散建设，存在数据来源异构、统计口径不一致、分析结果复用困难和决策依据滞后等问题，削弱了人才培养方案动态优化与就业治理的支撑能力。围绕高校“需求—招生—培养—就业—监测”全链路治理需求，本文设计并实现以关系型数据库为核心底座，融合数据治理、智能分析和可视化展示的一体化平台。",
    37: "围绕高校人才培养业务需求，本文确立平台的数据来源、指标体系和角色边界，构建覆盖数据接入、标准映射、质量清洗、标签归一、安全治理和分层入库的数据治理流程。在数据组织层面，系统以学生、企业、学业和就业等核心对象为中心建立关系型数据库模型，并面向需求预测、招生匹配、培养优化、就业推荐和动态监测等场景设计结果数据表，实现基础数据、分析特征和业务结果的分层存储与统一管理。在分析方法层面，平台结合时间趋势分析、匹配计算、关联规则挖掘和画像推荐等方法，对人才需求变化、专业生源适配、课程培养效果和学生就业方向进行综合分析，将结果展示扩展为面向决策支持的闭环分析过程。",
    38: "系统采用浏览器/服务器架构。后端承担数据查询、权限认证、审计留痕、结果封装和专报生成任务；前端面向教师端、政府端和公众端提供差异化可视化页面；数据库统一承载基础业务数据、安全数据和分析结果数据。实验与测试基于十万级脱敏仿真样本，从数据质量、功能完整性、接口连通性、结果解释性和安全审计等方面验证系统可运行性。结果表明，该平台能够支撑多源异构数据的统一组织、核心指标的可视化表达和分析结果的可解释输出，可为高校专业结构优化、招生策略调整、培养方案修订和就业指导提供可复用的数据支撑路径。",
    44: "The expansion of higher education and the restructuring of regional industries have intensified the mismatch between university talent supply and labor-market demand. In many universities, enrollment, academic affairs, curriculum training, and employment services are supported by independently developed information systems, which leads to heterogeneous data sources, inconsistent statistical definitions, limited reuse of analytical results, and delayed evidence for governance decisions. To address these issues, this thesis designs and implements an integrated platform for the full-chain governance scenario of demand, enrollment, training, employment, and monitoring. The platform uses a relational database as its core data foundation and integrates data governance, analytical modeling, and visual presentation.",
    45: "The study first defines the platform data sources, indicator system, and role-specific requirements according to the operational needs of university talent training and employment governance. It then establishes a data governance workflow that covers data ingestion, standard mapping, quality cleaning, label normalization, security control, and layered storage. A relational database model is constructed around core entities including students, enterprises, academic records, and employment records. Application-oriented result tables are further designed for predictive modeling of talent demand, enrollment matching, training optimization, employment recommendation, and dynamic monitoring. On this basis, time-trend analysis, matching computation, association-rule mining, and profile-based recommendation are combined to analyze changes in talent demand, enrollment suitability, curriculum-training effectiveness, and student employment orientation.",
    46: "The platform adopts a Browser/Server architecture. The back end provides data query, authentication, audit logging, result encapsulation, and report generation, whereas the front end offers role-specific visualization pages for teachers, government administrators, and public users. The database stores basic business data, security data, and analytical result data in a unified schema. Experiments are conducted on a de-identified simulated dataset with more than 100,000 records, covering data quality, functional completeness, interface connectivity, result interpretability, and security auditing. The results indicate that the proposed platform supports unified organization of multi-source heterogeneous data, visual representation of core indicators, and interpretable analytical output, thereby providing a reusable data-support path for major-structure optimization, enrollment-strategy adjustment, curriculum revision, and employment guidance in universities.",
    67: "本章阐明研究背景、研究意义、研究现状、研究内容、技术路线和全文组织结构。研究对象聚焦高校人才供需错配、校内数据孤岛和就业治理数字化转型等问题，提出面向“需求—招生—培养—就业—监测”链路的一体化平台设计与实现方案。",
    69: "高等教育大众化进程与产业结构调整共同推动高校人才培养模式转型。高校毕业生就业压力与企业高质量岗位空缺并存，表明人才供给结构与产业需求之间仍存在匹配偏差。高校的专业设置、招生规模、课程体系和就业服务质量直接影响人才供给质量。然而，招生、教务、培养和就业等业务系统通常由不同部门独立建设，数据标准、统计口径和业务流程缺乏统一规划，造成数据贯通困难、结果复用不足和决策依据滞后。数据孤岛削弱了高校对外部市场需求变化的感知能力，也限制了就业结果向招生策略、课程体系和培养方案调整的反馈。",
    70: "《中国教育现代化2035》提出加快信息化时代教育变革，高校需要进一步推进智慧校园和教育治理数字化建设。高等教育系统与劳动力市场之间存在由专业结构、招生计划、课程体系、实践训练、企业需求、产业政策和区域经济共同作用形成的复杂关系。单独观察招生规模、就业率或薪资水平，难以揭示市场需求变化、培养过程调整和就业质量提升之间的传导机制。因此，覆盖“需求—招生—培养—就业—监测”全过程的数据分析平台，是高校业务系统整合和教育治理过程化的重要基础。",
    71: "本文围绕高校人才培养与就业治理场景，设计并实现“需求—招生—培养—就业—监测”一体化大数据分析与可视化平台。平台整合外部招聘市场需求数据与校内招生、培养、就业数据，形成统一的数据治理、关系建模、分析计算和可视化表达链路。需求预测模块识别外部人才市场变化趋势，招生匹配模块辅助生源结构和专业布局优化，培养方案优化模块挖掘课程体系与就业结果之间的关联关系，就业推荐模块提高学生与岗位之间的匹配效率，动态监测与智能专报模块将分析结果反馈给管理者，支撑持续改进。",
    72: "本研究的意义体现在两方面。第一，统一数据底座和分析结果表设计能够提升高校招生、培养和就业数据的组织效率与复用能力，为数据分散、口径不一和分析滞后等问题提供系统化实现路径。第二，市场需求感知、培养过程优化和就业结果反馈被纳入同一平台后，高校能够对人才培养全过程进行动态监测和辅助决策，为专业结构优化、招生策略调整、课程体系修订和就业指导服务提供数据支撑。",
    74: "围绕高校人才培养、就业治理和教育数据应用，既有研究与工程实践主要集中在高校就业信息化、数据治理与数据库建模、可视化和智能决策支持三个方向。下文按上述方向梳理研究基础，并进一步界定本文平台的工程定位。",
    77: "现有系统的不足主要体现在分析链路不连续。部分高校系统侧重就业结果统计，较少将外部岗位需求、招生计划和培养课程纳入同一数据链路；部分系统能够展示就业率和薪资等指标，但缺少面向未来需求变化的预测、面向专业建设的规则证据以及面向多角色的权限隔离。本文平台以“需求—招生—培养—就业—监测”为主线，将就业结果分析前移至需求感知和培养调整环节。",
    80: "在数据库建模方面，数据仓库和维度建模提供了面向分析场景的常用方法[3]。本项目以原型系统为定位，在MySQL中采用DIM维度层、FACT事实层、ADS应用结果层和SYS安全审计层进行分层设计。该设计支持学校、专业、行业和岗位等编码在不同模块中复用，并使需求预测、招生匹配、培养优化和就业推荐结果能够稳定供接口查询。",
    83: "结合项目实现，平台建设的关键难点包括多源异构数据标准化、专业与岗位标签映射、数据质量控制、ADS结果表复用、教师端与政府端的数据范围隔离、公众端脱敏展示、统一登录鉴权和审计日志追踪。这些问题构成平台工程实现的主要约束。",
    85: "本文以高校“需求—招生—培养—就业—监测”一体化治理为研究对象，面向教育数字化转型和人才供需结构性矛盾，构建可复用的数据分析与可视化平台。平台以上海高校业务场景为应用背景，通过整合校内培养数据与校外产业需求信息，为高校学科建设、专业结构调整和就业服务优化提供量化依据，并为教育主管部门开展资源布局和就业治理提供辅助决策支撑。",
    86: "在技术实现层面，平台包括数据生成与接入、分层存储、特征加工、算法分析、接口服务和多角色可视化展示等模块。底层数据架构采用MySQL分层数据库承载维度表、事实表、应用结果表和安全审计表；数据处理脚本完成清洗、脱敏、标准化和聚合计算；需求预测模块基于历史岗位需求序列、季节因子、政策热度和专业强度生成未来需求结果；培养优化模块利用FP-Growth算法挖掘课程、技能、岗位和就业结果之间的关联规则；就业推荐模块基于学生画像与岗位画像计算余弦相似度；后端采用Flask提供接口服务，前端通过React、Ant Design和ECharts实现多角色可视化页面。",
    87: "本研究的目标包括三项：构建贯通需求、招生、培养、就业和监测的数据链路；实现以关系型数据库结果表为核心的B/S架构平台；验证核心接口、数据结果和多角色页面的可运行性。在算法评价层面，需求预测、培养规则和就业推荐结果需具备可查询、可解释和可复用特征；在系统评价层面，平台需满足基础连通性、权限隔离和结果展示要求。",
    88: "本文通过数据治理、关系型数据库建模、算法结果表设计和Web可视化实现，将原始业务数据转换为可供管理决策使用的结构化指标与分析结果。",
    100: "本章围绕平台建设所需的关键理论与技术基础展开。系统涉及数据库建模、数据治理、特征加工、预测模型、推荐算法、关联规则挖掘、Web接口和可视化展示。本章按照“数据底座—计算模型—系统架构—安全治理”的逻辑说明相关技术及其在项目中的应用位置。",
    116: "LSTM是循环神经网络的改进结构，适合处理具有时间依赖关系的序列数据[6]。在人才需求预测场景中，岗位需求人数受毕业周期、招聘季节、产业政策和区域经济变化影响，表现出时间波动性。本文将预测目标定义为岗位需求人数或人才需求热度，薪资水平作为就业质量分析和推荐解释的辅助特征。",
    117: "令 Y = {y_1, y_2, …, y_n} 表示按时间顺序排列的真实岗位需求序列，Ŷ = {ŷ_1, ŷ_2, …, ŷ_n} 表示模型预测序列。其中，n 为时间序列的样本总数，i 为第 i 个时间节点，y_i 为第 i 个时间节点的真实需求值，ŷ_i 为第 i 个时间节点的预测值。本文采用的误差指标如式（2-1）至式（2-3）所示。",
    118: "MAE = (1/n) Σ_{i=1}^{n} |y_i − ŷ_i|",
    119: "RMSE = sqrt[(1/n) Σ_{i=1}^{n} (y_i − ŷ_i)^2]",
    120: "MAPE = (100%/n) Σ_{i=1}^{n} |(y_i − ŷ_i) / y_i|",
    124: "当前项目的scripts/LSTM-all.py采用时间序列预测思想与季节性回退策略结合的工程实现。脚本未直接依赖TensorFlow训练深度模型，而是基于历史需求序列、季节因子、政策热度和专业强度生成未来12个月预测结果，并写入ads_job_demand_forecast、ads_job_demand_forecast_eval和ads_job_demand_forecast_backtest。",
    131: "关联规则挖掘用于发现事务数据中不同项之间的共现关系。FP-Growth通过构建频繁模式树减少候选集生成开销，适合在课程、技能、岗位和就业结果等离散标签中挖掘高频组合[8]。令 D = {T_1, T_2, …, T_N} 表示事务数据库，N 为事务总数，T_j 为第 j 条事务；令 A 和 B 分别表示两个不相交项集，且 A ∩ B = ∅。规则 A ⇒ B 的支持度、置信度和提升度定义如式（2-4）至式（2-6）所示。",
    132: "support(A ⇒ B) = count(A ∪ B) / N",
    133: "confidence(A ⇒ B) = count(A ∪ B) / count(A)",
    134: "lift(A ⇒ B) = confidence(A ⇒ B) / support(B)",
    135: "其中，count(A ∪ B) 表示同时包含项集 A 与项集 B 的事务数量，count(A) 表示包含项集 A 的事务数量，support(B) 表示项集 B 在事务数据库中的支持度。当前项目中，scripts/FPgrowth-all.py读取fact_course_skill、fact_employment和dim_job_category等数据，生成ads_training_rules。该表保存antecedents、consequents、support、confidence、lift和evidence_score，为培养优化页面提供规则证据。",
    137: "余弦相似度用于衡量两个向量方向的一致程度，适合描述学生画像与岗位画像之间的匹配关系。令 a = (a_1, a_2, …, a_m) 表示学生画像向量，b = (b_1, b_2, …, b_m) 表示岗位画像向量。其中，m 为画像特征维度总数，a_j 为学生在第 j 个特征维度上的取值，b_j 为岗位在第 j 个特征维度上的取值。余弦相似度定义如式（2-7）所示。",
    138: "cos(a, b) = (a · b) / (||a||₂ ||b||₂)\n= [Σ_{j=1}^{m} a_j b_j] / [sqrt(Σ_{j=1}^{m} a_j^2) sqrt(Σ_{j=1}^{m} b_j^2)]",
    139: "当前项目的就业推荐模块综合学生专业、技能标签、求职意向、岗位需求强度、历史就业质量和企业属性计算相似度，并输出similarity_score、confidence_level、salary_reference和recommendation_reason等字段。相似度越高，表示学生画像与岗位画像在特征空间中的方向越接近。",
    151: "平台面向高校人才培养与就业治理场景，目标是将外部产业需求、招生入口、培养过程、就业结果和动态监测组织为连续业务链条。教师或专业负责人需要掌握本校专业就业表现、课程能力短板和岗位匹配方向；教育主管部门需要掌握区域高校专业布局、就业质量差异和重点产业人才供给；公众用户需要获得公开、可理解的院校与专业就业信息。",
    152: "业务场景的关键矛盾在于信息链条断裂。外部招聘市场的岗位需求变化难以及时反馈至专业招生和培养方案，学生就业结果也难以反向解释课程体系与专业结构的有效性。平台通过统一数据库和结果表体系串联上述数据节点，使各模块共享同一底层事实。",
    153: "系统采用高仿真样本验证链路可运行性。样本覆盖学生基础信息、学业事实、就业事实和企业维度属性，并通过学校层级、专业类别、行业标签、企业规模和薪资水平等变量构造教育就业场景。",
    157: "平台设置教师端、政府端、公众端和管理端等角色。教师端面向专业负责人、学院管理人员和就业指导教师，主要使用需求预测、招生匹配、培养方案优化、就业推荐和规则证据库等模块。政府端面向教育主管部门或区域治理部门，重点使用宏观监测、学校对标、需求预测、智能专报和政策辅助分析等模块。公众端面向学生、家长和社会用户，提供院校对比、专业就业质量、热门行业去向和公开就业趋势。",
    158: "不同角色共享同一数据底座，但页面权限和指标粒度不同。教师端可查看本校及相关专业的细粒度分析结果，政府端可查看跨学校汇总比较，公众端仅展示脱敏后的公开信息。管理端负责账号、角色、审计和系统配置。",
    162: "系统非功能需求包括数据一致性、查询性能、安全性、可扩展性、可维护性和可解释性。数据一致性要求学校、专业、行业、地区和岗位等字段采用统一编码与口径。查询性能要求前端高频页面优先读取结果表，避免每次访问重新聚合明细数据。安全性要求密码哈希存储、角色权限控制和操作日志记录。",
    167: "平台数据对象覆盖学校、专业、企业、岗位需求、毕业生、就业结果、招生计划、课程技能、政策信号、安全账号和审计日志等类型。根据当前项目实现，学校与专业数据进入dim_school、dim_major_catalog和bridge_school_major；企业、行业与岗位类别进入dim_enterprise、dim_industry和dim_job_category；岗位、毕业、就业、招生、课程和政策等业务事实进入FACT层；账号、角色和审计信息进入SYS层。",
    171: "系统采用B/S架构，整体划分为数据源层、数据治理层、数据库层、算法分析层、接口服务层和应用展示层。数据源层包括学生、学业、就业、企业、招聘和政策等数据；数据治理层负责清洗、标准化、标签归一和脱敏；数据库层保存基础表、事实表、应用结果表和安全审计表；算法分析层负责预测、匹配、规则挖掘和推荐计算；接口服务层基于Flask提供统一REST服务；应用展示层基于React和ECharts形成多角色页面。",
    172: "该架构以数据库结果表作为上下层之间的稳定契约。算法脚本独立运行并写入结果表，后端服务按接口读取结果表并执行权限校验，前端页面无需感知算法内部细节。该设计降低模块耦合，并提升问题定位效率。",
    178: "图3-2表明，数据源层提供原始业务对象，数据治理层完成清洗、标签归一和脱敏处理，数据库层以DIM、FACT、ADS和SYS保存稳定数据，算法分析层批量写入结果表，接口服务层负责角色校验和JSON封装，应用展示层面向教师端、政府端和公众端进行差异化呈现。",
    180: "系统详细设计需要说明业务模块协同方式、数据库承接算法结果的方式以及用户界面对多角色使用场景的支撑方式。参照系统分层架构，本文将详细设计划分为功能模块设计、数据库设计和用户界面设计三部分。设计逻辑是以数据库表结构作为稳定契约，以离线算法脚本形成可复现结果，以后端接口完成权限校验和结果封装，并通过前端页面向教师端、政府端和公众端提供差异化服务。",
    182: "平台功能设计围绕“需求—招生—培养—就业—监测”业务闭环展开。系统不局限于展示就业率、薪资等结果指标，而是将外部岗位需求、招生计划、培养课程、毕业生画像和就业去向纳入同一数据链路，使需求变化传导至招生匹配和培养优化，并通过学生就业结果反向验证专业结构与课程体系的适配程度。",
    183: "从工程实现角度看，系统功能分为八个模块。数据治理与入库模块负责生成或接入多源数据，并完成清洗、脱敏、标准映射和分层入库；岗位需求预测模块依据岗位月度需求、政策热度、专业强度和薪资水平预测未来需求；招生匹配模块综合就业质量、需求增长、政策热度和专业基础给出专业匹配分；培养优化模块依据课程、技能、岗位和就业结果之间的关联规则提出调整建议；就业推荐模块根据学生画像和岗位画像计算匹配度；动态监测模块面向政府端和公众端展示学校对标、产业流向和风险预警；智能专报模块将数据库指标组织为可解释文本；权限审计模块负责账号、角色、访问范围和操作留痕。其结构如图3-3所示。",
    186: "图3-3表明，平台功能模块并非独立页面集合，而是以数据治理和ADS结果表为共同基础形成的业务闭环。需求预测结果进入招生匹配、培养优化和就业推荐模块；培养规则和专业优化建议为教师端专业建设提供解释依据；政府端和公众端在权限过滤后读取汇总结果，避免明细数据越权展示。",
    188: "在模块调用关系上，系统采用离线计算与在线服务结合的方式。数据生成、特征加工、模型预测、规则挖掘和推荐计算由脚本批量执行并写入结果表；在线服务负责身份认证、参数过滤、SQL查询、结果格式化和异常处理。该设计避免前端访问时重复执行复杂计算，并便于通过算法链路日志追踪每个阶段的输入表、输出表、运行状态和执行耗时。",
    193: "图3-4说明了数据库设计的关键原则：第一，基础业务数据与分析结果数据分离，避免模型结果反向污染原始事实；第二，学校、专业、行业和岗位等编码在各层保持一致，使跨模块追溯成为可能；第三，前端页面优先读取ADS应用层，保证展示性能和接口稳定性；第四，安全审计表独立于业务表，便于后续扩展账号策略、操作追踪和风险审计功能。",
    197: "应用结果表采用“按场景冗余关键展示字段”的方式设计。例如，ads_job_demand_forecast不仅保存school_id、major_code等编码，还保存school_name、major_name、industry_name和job_category_name，便于前端直接展示；ads_job_recommendation保存推荐岗位、企业、相似度、置信等级和推荐理由，避免页面再进行复杂拼接；ads_major_optimization保存建议类型、优先级、证据分、原因说明和相关指标，使教师端能够解释专业扩招、稳招、缩招或加强实践培养等建议的依据。",
    203: "用户界面采用B/S架构，前端使用React、Ant Design和ECharts实现。系统入口为登录页，用户登录后根据角色进入不同工作台。教师端默认进入上海大学口径的专业建设与就业治理工作台，政府端进入区域高校治理驾驶舱，公众端进入公开展示平台。不同端口共享组件体系和数据服务层，但菜单结构、指标粒度和可访问接口不同。",
    205: "在交互设计上，教师端页面强调可钻取和可解释。需求预测页面使用折线图展示未来12个月岗位需求人数，并通过图例区分专业和岗位组合；招生匹配页面使用KPI卡片、画像条形图和专业列表展示匹配分；培养优化页面使用气泡图和建议卡片展示优先级、建议类型和原因；就业推荐页面提供学生ID查询、Top-K推荐列表和推荐理由。政府端页面强调对比和预警，通过学校榜单、产业分布、区域预警和专报文本支撑宏观治理。公众端减少内部管理字段，仅保留适合公开展示的院校对比和专业趋势。",
    305: "接口测试采用scripts/check_api_contract.py生成的报告。受限于本地测试环境的硬件条件与测试工具记录范围，本文未开展高并发场景下的毫秒级响应延迟测试，仅验证核心接口的连通性、返回结构一致性和页面数据可用性。测试结果显示，教师端、政府端和公众端核心接口均可在本地环境下正常返回，页面未出现空数据和全0指标风险。",
    312: "本章从测试数据、数据质量、功能、接口、算法结果和权限审计等方面验证系统可运行性。测试结果表明，平台能够支撑核心业务页面展示和算法结果查询。受限于脱敏仿真数据和本地测试环境，真实授权数据接入后仍需重新校准模型参数，并补充浏览器自动化测试、接口压测和安全测试。",
    319: "6.3 不足与展望",
    320: "当前系统仍存在以下局限。首先，实验数据主要为脱敏仿真样本，真实授权数据接入后需要重新进行字段映射、质量校验和模型参数校准。其次，需求预测采用时间序列预测思想与季节性回退结合的工程实现，尚未在真实长周期序列上完成完整深度学习训练与泛化验证。再次，当前测试主要依赖接口合约、静态扫描和构建检查，尚未覆盖高并发压测、浏览器自动化回归测试和更严格的安全测试。",
}


def replace_paragraph_text(paragraph, text):
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def main():
    doc = Document(SRC)
    for idx, text in replacements.items():
        replace_paragraph_text(doc.paragraphs[idx], text)

    # LaTeX fragments render poorly under the thesis template's distributed
    # Chinese alignment, so formula-related paragraphs get explicit alignment.
    explanation_paragraphs = [117, 131, 135, 137]
    equation_paragraphs = [118, 119, 120, 132, 133, 134, 138]
    for idx in explanation_paragraphs:
        p = doc.paragraphs[idx]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for idx in equation_paragraphs:
        p = doc.paragraphs[idx]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(9.5 if idx == 138 else 10.5)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
