from __future__ import annotations

import json
import re
import shutil
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


SRC = Path(r"D:\VsProject\CAC_DB\bise_working.docx")
BACKUP = Path(r"D:\VsProject\CAC_DB\bise_original_backup.docx")
OUT = Path(r"D:\VsProject\CAC_DB\毕设_图表题注与表格格式统一.docx")
REPORT = Path(r"D:\VsProject\CAC_DB\caption_table_report.json")

CN_NUM = {
    "一": 1,
    "二": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
    "十": 10,
}

CHAPTER_RE = re.compile(r"^第([一二三四五六七八九十0-9]+)章")
FIG_CAP_RE = re.compile(r"^\s*图\s*(?:[0-9一二三四五六七八九十]+(?:\s*[-.．－—]\s*[0-9]+)?|[A-Za-z]+)\s*[:：、]?\s*(.*)$")
TBL_CAP_RE = re.compile(r"^\s*表\s*(?:[0-9一二三四五六七八九十]+(?:\s*[-.．－—]\s*[0-9]+)?|[A-Za-z]+)\s*[:：、]?\s*(.*)$")


def cn_to_int(text: str) -> int | None:
    if text.isdigit():
        return int(text)
    if text in CN_NUM:
        return CN_NUM[text]
    if text.startswith("十"):
        return 10 + CN_NUM.get(text[1:], 0)
    if "十" in text:
        left, right = text.split("十", 1)
        return CN_NUM.get(left, 1) * 10 + CN_NUM.get(right, 0)
    return None


def element_text(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def para_style_id(el) -> str:
    p_pr = el.find(qn("w:pPr"))
    if p_pr is None:
        return ""
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        return ""
    return p_style.get(qn("w:val")) or ""


def has_picture(el) -> bool:
    return bool(list(el.iter(qn("w:drawing"))) or list(el.iter(qn("w:pict"))))


def is_paragraph(el) -> bool:
    return el.tag == qn("w:p")


def is_table(el) -> bool:
    return el.tag == qn("w:tbl")


def paragraph_from_el(el, doc: Document) -> Paragraph:
    return Paragraph(el, doc)


def table_from_el(el, doc: Document) -> Table:
    return Table(el, doc)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    format_run(run)


def format_run(run) -> None:
    run.bold = False
    run.font.name = "宋体"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "宋体")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")


def format_caption(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1
    for run in paragraph.runs:
        format_run(run)


def add_paragraph_after(el, text: str, doc: Document) -> Paragraph:
    new_p = OxmlElement("w:p")
    el.addnext(new_p)
    paragraph = paragraph_from_el(new_p, doc)
    set_paragraph_text(paragraph, text)
    format_caption(paragraph)
    return paragraph


def add_paragraph_before(el, text: str, doc: Document) -> Paragraph:
    new_p = OxmlElement("w:p")
    el.addprevious(new_p)
    paragraph = paragraph_from_el(new_p, doc)
    set_paragraph_text(paragraph, text)
    format_caption(paragraph)
    return paragraph


def strip_caption_title(text: str, kind: str) -> str:
    regex = FIG_CAP_RE if kind == "fig" else TBL_CAP_RE
    match = regex.match(text)
    if not match:
        return ""
    title = (match.group(1) or "").strip()
    title = re.sub(r"^[\s:：、，,.-]+", "", title).strip()
    return title


def first_row_text(table_el) -> str:
    rows = table_el.findall(qn("w:tr"))
    if not rows:
        return ""
    cells = rows[0].findall(qn("w:tc"))
    parts = [element_text(cell) for cell in cells]
    return "、".join(p for p in parts if p)


def table_fallback_title(table_el) -> str:
    header = first_row_text(table_el)
    if header:
        return "、".join(header.split("、")[:4]) + "表"
    return "待补充表名"


def find_next_caption(body, idx: int, regex: re.Pattern, limit: int = 5) -> int | None:
    for j in range(idx + 1, min(len(body), idx + limit + 1)):
        if is_table(body[j]):
            return None
        if is_paragraph(body[j]):
            text = element_text(body[j])
            if text and regex.match(text):
                return j
            if text and has_picture(body[j]):
                return None
    return None


def find_prev_caption(body, idx: int, regex: re.Pattern, limit: int = 5) -> int | None:
    for j in range(idx - 1, max(-1, idx - limit - 1), -1):
        if is_table(body[j]):
            return None
        if is_paragraph(body[j]):
            text = element_text(body[j])
            if text and regex.match(text):
                return j
            if text and has_picture(body[j]):
                return None
    return None


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_table_borders(table: Table) -> None:
    tbl_pr = ensure_child(table._tbl, "w:tblPr")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def set_cell_white(cell) -> None:
    tc_pr = ensure_child(cell._tc, "w:tcPr")
    for shd in list(tc_pr.findall(qn("w:shd"))):
        tc_pr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFFFF")
    tc_pr.append(shd)


def set_run_black(run) -> None:
    r_pr = run._element.get_or_add_rPr()
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    color.attrib.pop(qn("w:themeColor"), None)
    color.attrib.pop(qn("w:themeTint"), None)
    color.attrib.pop(qn("w:themeShade"), None)
    color.set(qn("w:val"), "000000")


def normalize_table(table: Table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            set_cell_white(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_black(run)


def collect_chapters(body) -> dict[int, int]:
    chapters = {}
    current = 0
    for i, el in enumerate(body):
        if not is_paragraph(el):
            continue
        text = element_text(el)
        style = para_style_id(el)
        match = CHAPTER_RE.match(text)
        if match and not style.startswith("TOC"):
            num = cn_to_int(match.group(1))
            if num:
                current = num
        chapters[i] = current
    return chapters


def current_chapter_for(chapters: dict[int, int], idx: int) -> int:
    keys = [k for k in chapters if k <= idx]
    if not keys:
        return 0
    return chapters[max(keys)]


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(SRC, BACKUP)

    doc = Document(str(SRC))
    body = list(doc.element.body.iterchildren())

    start_idx = None
    ref_idx = len(body)
    for i, el in enumerate(body):
        if not is_paragraph(el):
            continue
        text = element_text(el)
        style = para_style_id(el)
        if start_idx is None and CHAPTER_RE.match(text) and not style.startswith("TOC"):
            start_idx = i
        if start_idx is not None and text.startswith("参考文献"):
            ref_idx = i
            break
        if start_idx is not None and text.startswith("（图标题中文"):
            ref_idx = min(ref_idx, i)
            break
    start_idx = start_idx or 0

    chapters = collect_chapters(body)
    fig_counts = defaultdict(int)
    tbl_counts = defaultdict(int)
    report = {
        "body_images": 0,
        "figure_captions_normalized_or_added": 0,
        "body_tables": 0,
        "table_captions_normalized_or_added": 0,
        "tables_formatted": 0,
        "skipped_cover_tables": 0,
        "skipped_template_examples": 0,
        "skipped_decorative_images": 0,
        "pending_figure_titles": [],
        "pending_table_titles": [],
        "figure_numbers": [],
        "table_numbers": [],
    }

    formal_table_indices = []

    for idx, el in enumerate(body):
        chapter = current_chapter_for(chapters, idx)
        in_body = start_idx <= idx < ref_idx and chapter >= 1

        if has_picture(el):
            if not in_body:
                if idx < start_idx:
                    report["skipped_decorative_images"] += 1
                else:
                    report["skipped_template_examples"] += 1
                continue
            report["body_images"] += 1
            fig_counts[chapter] += 1
            number = f"{chapter}-{fig_counts[chapter]}"
            cap_idx = find_next_caption(body, idx, FIG_CAP_RE)
            title = ""
            if cap_idx is not None:
                title = strip_caption_title(element_text(body[cap_idx]), "fig")
            if not title:
                title = "待补充图名"
                report["pending_figure_titles"].append(f"图{number}：正文元素 {idx}")
            text = f"图{number} {title}"
            if cap_idx is not None:
                paragraph = paragraph_from_el(body[cap_idx], doc)
                set_paragraph_text(paragraph, text)
                format_caption(paragraph)
            else:
                add_paragraph_after(el, text, doc)
                body = list(doc.element.body.iterchildren())
            report["figure_captions_normalized_or_added"] += 1
            report["figure_numbers"].append(f"图{number}")

        elif is_table(el):
            if not in_body:
                if idx < start_idx:
                    report["skipped_cover_tables"] += 1
                else:
                    report["skipped_template_examples"] += 1
                continue
            formal_table_indices.append(idx)
            report["body_tables"] += 1
            tbl_counts[chapter] += 1
            number = f"{chapter}-{tbl_counts[chapter]}"
            cap_idx = find_prev_caption(body, idx, TBL_CAP_RE)
            title = ""
            if cap_idx is not None:
                title = strip_caption_title(element_text(body[cap_idx]), "tbl")
            if not title:
                title = table_fallback_title(el)
                if title == "待补充表名":
                    report["pending_table_titles"].append(f"表{number}：正文元素 {idx}")
            text = f"表{number} {title}"
            if cap_idx is not None:
                paragraph = paragraph_from_el(body[cap_idx], doc)
                set_paragraph_text(paragraph, text)
                format_caption(paragraph)
            else:
                add_paragraph_before(el, text, doc)
                body = list(doc.element.body.iterchildren())
            normalize_table(table_from_el(el, doc))
            report["table_captions_normalized_or_added"] += 1
            report["tables_formatted"] += 1
            report["table_numbers"].append(f"表{number}")

    doc.save(str(OUT))

    qa = quality_check(OUT, start_idx, ref_idx)
    report["quality_check"] = qa
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


def quality_check(path: Path, start_idx: int, ref_idx: int) -> dict:
    doc = Document(str(path))
    body = list(doc.element.body.iterchildren())
    chapters = collect_chapters(body)
    duplicate_caption_pairs = []
    nonwhite_shading = []
    nonblack_table_text = []
    fig_by_chapter = defaultdict(list)
    tbl_by_chapter = defaultdict(list)

    for i, el in enumerate(body):
        if is_paragraph(el):
            text = element_text(el)
            m_fig = re.match(r"^图(\d+)-(\d+)\s+", text)
            if m_fig:
                fig_by_chapter[int(m_fig.group(1))].append(int(m_fig.group(2)))
            m_tbl = re.match(r"^表(\d+)-(\d+)\s+", text)
            if m_tbl:
                tbl_by_chapter[int(m_tbl.group(1))].append(int(m_tbl.group(2)))
                if i + 1 < len(body) and is_paragraph(body[i + 1]):
                    next_text = element_text(body[i + 1])
                    if re.match(r"^表\d+-\d+\s+", next_text):
                        duplicate_caption_pairs.append([i, i + 1])
            if re.match(r"^图\d+-\d+\s+", text):
                for j in range(i + 1, min(i + 3, len(body))):
                    if is_paragraph(body[j]) and re.match(r"^图\d+-\d+\s+", element_text(body[j])):
                        duplicate_caption_pairs.append([i, j])

        if is_table(el) and start_idx <= i < ref_idx and current_chapter_for(chapters, i) >= 1:
            for shd in el.iter(qn("w:shd")):
                fill = (shd.get(qn("w:fill")) or "").upper()
                if fill and fill not in {"FFFFFF", "AUTO"}:
                    nonwhite_shading.append(i)
            for color in el.iter(qn("w:color")):
                val = (color.get(qn("w:val")) or "").upper()
                if val and val not in {"000000", "AUTO"}:
                    nonblack_table_text.append(i)

    fig_sequence_ok = all(nums == list(range(1, len(nums) + 1)) for nums in fig_by_chapter.values())
    tbl_sequence_ok = all(nums == list(range(1, len(nums) + 1)) for nums in tbl_by_chapter.values())
    return {
        "reopen_ok": True,
        "duplicate_caption_pairs": duplicate_caption_pairs,
        "figure_sequence_ok": fig_sequence_ok,
        "table_sequence_ok": tbl_sequence_ok,
        "nonwhite_shading_table_elements": sorted(set(nonwhite_shading)),
        "nonblack_text_table_elements": sorted(set(nonblack_table_text)),
    }


if __name__ == "__main__":
    main()
