from __future__ import annotations

import copy
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE_SIZE = 10_112_216
SOURCE_DIR = Path(r"C:\Users\1\Desktop")
OUT_NAME = "最新毕设_仅图表格式调整.docx"
REPORT_NAME = "图表格式调整报告.txt"
TEMP_OUT = "latest_bise_fig_table_format_tmp.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CAPTION_RE = re.compile(r"^\s*([图表])\s*\d+[\-.]\d+\s+")


def find_source() -> Path:
    local = Path("最新毕设.docx")
    if local.exists():
        return local
    matches = [p for p in SOURCE_DIR.glob("*.docx") if p.stat().st_size == SOURCE_SIZE]
    if matches:
        return matches[0]
    fallback = SOURCE_DIR / "最新毕设.docx"
    if fallback.exists():
        return fallback
    raise FileNotFoundError("未找到 最新毕设.docx")


def iter_body_blocks(doc: Document):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "p", Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield "t", Table(child, doc)


def has_field_or_bookmark(paragraph: Paragraph) -> bool:
    return bool(
        paragraph._p.xpath(
            ".//w:fldChar|.//w:instrText|.//w:fldSimple|.//w:bookmarkStart|.//w:bookmarkEnd"
        )
    )


def set_run_fonts(run, east_asia: str, ascii_font: str, size_pt: float, color: str | None = None) -> None:
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), ascii_font)
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def is_ascii_caption_char(ch: str) -> bool:
    return ord(ch) < 128


def split_caption_text(text: str) -> list[tuple[str, bool]]:
    if not text:
        return []
    groups: list[tuple[str, bool]] = []
    buf = text[0]
    cur = is_ascii_caption_char(text[0])
    for ch in text[1:]:
        flag = is_ascii_caption_char(ch)
        if flag == cur:
            buf += ch
        else:
            groups.append((buf, cur))
            buf = ch
            cur = flag
    groups.append((buf, cur))
    return groups


def clear_text_runs_keep_ppr(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag == qn("w:r") or child.tag == qn("w:hyperlink"):
            p.remove(child)


def append_styled_run(paragraph: Paragraph, text: str, ascii_group: bool) -> None:
    run = paragraph.add_run(text)
    set_run_fonts(run, "黑体", "Times New Roman", 12, "000000")
    run.bold = bool(ascii_group)


def style_caption_paragraph(paragraph: Paragraph, unsafe: list[str], para_idx: int) -> str:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1

    text = paragraph.text
    if has_field_or_bookmark(paragraph):
        for run in paragraph.runs:
            set_run_fonts(run, "黑体", "Times New Roman", 12, "000000")
            if run.text and all(is_ascii_caption_char(ch) for ch in run.text):
                run.bold = True
        unsafe.append(f"段落{para_idx + 1}：{text}")
        return text

    clear_text_runs_keep_ppr(paragraph)
    for segment, ascii_group in split_caption_text(text):
        append_styled_run(paragraph, segment, ascii_group)
    return text


def ensure_tbl_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def set_cell_shading_white(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFFFF")


def style_table(table: Table) -> None:
    ensure_tbl_borders(table)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading_white(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_fonts(run, "宋体", "Times New Roman", 10.5, "000000")


def body_text_snapshot(doc: Document) -> tuple[list[str], list[list[list[str]]]]:
    paras = [p.text for p in doc.paragraphs]
    tables = []
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            table_rows.append([cell.text for cell in row.cells])
        tables.append(table_rows)
    return paras, tables


def count_fields(docx_path: Path) -> tuple[int, int, int, list[str]]:
    instr_count = 0
    fld_char_count = 0
    fld_simple_count = 0
    instr_texts: list[str] = []
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml")
    root = __import__("lxml").etree.fromstring(xml)
    ns = {"w": W_NS}
    for el in root.xpath(".//w:instrText", namespaces=ns):
        instr_count += 1
        if el.text:
            instr_texts.append(el.text.strip())
    fld_char_count = len(root.xpath(".//w:fldChar", namespaces=ns))
    fld_simple_count = len(root.xpath(".//w:fldSimple", namespaces=ns))
    return instr_count, fld_char_count, fld_simple_count, instr_texts


def count_captions(doc: Document) -> tuple[list[tuple[int, str]], list[tuple[int, str]]]:
    figures: list[tuple[int, str]] = []
    tables: list[tuple[int, str]] = []
    in_main = False
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("第1章"):
            in_main = True
        if text.startswith("参考文献") or text.startswith("致  谢") or text.startswith("致谢"):
            in_main = False
        if not in_main:
            continue
        match = CAPTION_RE.match(text)
        if not match:
            continue
        if match.group(1) == "图":
            figures.append((idx, text))
        else:
            tables.append((idx, text))
    return figures, tables


def main() -> None:
    src = find_source()
    before_doc = Document(str(src))
    before_paras, before_tables_text = body_text_snapshot(before_doc)
    before_figs, before_tabs = count_captions(before_doc)
    before_table_count = len(before_doc.tables)
    instr_before, fld_before, simple_before, instr_texts = count_fields(src)

    doc = Document(str(src))
    fig_count = 0
    tab_caption_count = 0
    table_count = 0
    unsafe_captions: list[str] = []
    skipped_front_tables = 0

    in_main = False
    para_idx = -1
    table_idx = -1
    for kind, block in iter_body_blocks(doc):
        if kind == "p":
            para_idx += 1
            text = block.text.strip()
            if text.startswith("第1章"):
                in_main = True
            if text.startswith("参考文献") or text.startswith("致  谢") or text.startswith("致谢"):
                in_main = False
            if not in_main:
                continue
            match = CAPTION_RE.match(text)
            if not match:
                continue
            style_caption_paragraph(block, unsafe_captions, para_idx)
            if match.group(1) == "图":
                fig_count += 1
            else:
                tab_caption_count += 1
        else:
            table_idx += 1
            if in_main:
                style_table(block)
                table_count += 1
            else:
                skipped_front_tables += 1

    tmp = Path(TEMP_OUT)
    out = Path(OUT_NAME)
    doc.save(tmp)
    shutil.copy2(tmp, out)

    check_doc = Document(str(out))
    after_paras, after_tables_text = body_text_snapshot(check_doc)
    after_figs, after_tabs = count_captions(check_doc)
    instr_after, fld_after, simple_after, _ = count_fields(out)

    text_unchanged = before_paras == after_paras and before_tables_text == after_tables_text
    captions_unchanged = before_figs == after_figs and before_tabs == after_tabs
    tables_unchanged = before_table_count == len(check_doc.tables)
    fields_unchanged = (instr_before, fld_before, simple_before) == (
        instr_after,
        fld_after,
        simple_after,
    )
    has_cross_ref_field = any(
        re.search(r"\b(REF|PAGEREF|SEQ)\b", item, re.IGNORECASE) for item in instr_texts
    )
    has_any_field = bool(instr_before or fld_before or simple_before)

    report_lines = [
        "图表格式调整报告",
        "",
        f"源文件：{src.resolve()}",
        f"输出文件：{out.resolve()}",
        "",
        f"共识别并处理图标题：{fig_count}个。",
        f"共识别并处理表标题：{tab_caption_count}个。",
        f"共处理正文表格：{table_count}个。",
        f"已跳过封面/声明/目录区表格：{skipped_front_tables}个；参考文献和致谢区域未处理。",
        f"是否检测到交叉引用相关域（REF/PAGEREF/SEQ）：{'是' if has_cross_ref_field else '否'}。",
        f"是否检测到文档域代码：{'是' if has_any_field else '否'}；处理过程中未更新目录和交叉引用。",
        f"是否保持图表编号和题注文字不变：{'是' if captions_unchanged else '否'}。",
        f"是否保持正文段落和表格文字内容不变：{'是' if text_unchanged else '否'}。",
        f"是否未新增或删除表格：{'是' if tables_unchanged else '否'}。",
        f"域代码数量是否保持不变：{'是' if fields_unchanged else '否'}。",
    ]
    if unsafe_captions:
        report_lines.append("")
        report_lines.append("无法安全拆分的题注如下，未重写文本，仅做已有 run 的安全样式调整：")
        report_lines.extend(f"- {item}" for item in unsafe_captions)
    else:
        report_lines.append("")
        report_lines.append("未发现无法安全处理的题注。")
    report_lines.extend(
        [
            "",
            "质量检查：生成后的 docx 已重新打开读取；图标题、表标题、表格数量与文字快照均已比对。",
        ]
    )

    Path(REPORT_NAME).write_text("\n".join(report_lines), encoding="utf-8")


if __name__ == "__main__":
    main()
