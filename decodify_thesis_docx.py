from __future__ import annotations

import re
import shutil
from collections import Counter
from pathlib import Path

from docx import Document


SOURCE_SIZE = 9_802_029
OUT_NAME = "毕设_去代码化表达润色版.docx"
BACKUP_NAME = "毕设_图表题注与表格格式统一_去代码化前备份.docx"
REPORT_NAME = "去代码化表达修改报告.txt"


WATCH_TERMS = [
    "dim_school",
    "dim_major_catalog",
    "bridge_school_major",
    "dim_enterprise",
    "dim_industry",
    "dim_job_category",
    "fact_job_posting",
    "fact_graduate",
    "fact_employment",
    "fact_enrollment_plan",
    "fact_course_skill",
    "fact_policy_signal",
    "ads_job_demand_features",
    "ads_job_demand_forecast",
    "ads_job_recommendation",
    "ads_major_optimization",
    "ads_algorithm_chain_log",
    "ads_employment_summary",
    "ads_enrollment_matching",
    "ads_training_rules",
    "sys_user_account",
    "sys_audit_log",
    "dim_student",
    "dim_company",
    "fact_academic",
    "school_id",
    "major_code",
    "industry_id",
    "job_category_id",
    "user_id",
    "posting_id",
    "graduate_id",
    "employment_id",
    "month",
    "forecast_month",
    "demand_count_sum",
    "posting_count",
    "avg_salary",
    "policy_heat",
    "employment_rate",
    "match_score",
    "skill_gap_score",
    "enrollment_pressure",
    "school_major_strength_score",
    "predicted_demand_count",
    "lower_bound",
    "upper_bound",
    "demand_growth_rate",
    "demand_level",
    "mape",
    "track_rank",
    "similarity_score",
    "confidence_level",
    "salary_reference",
    "recommendation_reason",
    "password_hash",
    "hash_algo",
    "app.py",
    "backend/app.py",
    "dataService.js",
    "dataService",
    "platform_data_factory.py",
    "create_tables.py",
    "PutData.py",
    "Spark-all.py",
    "FPgrowth-all.py",
    "CF-all.py",
    "run_full_pipeline_check.py",
    "HTTP 200",
    "code=0",
    "code、message、data",
    "code、message和data",
    "/api",
    "API合约检查报告",
    "前端静态扫描",
    "依赖安装",
    "验收脚本执行15个步骤",
    "DIM",
    "FACT",
    "ADS",
    "SYS",
]

RESIDUAL_TERMS = [
    "dim_",
    "fact_",
    "ads_",
    "sys_",
    "school_id",
    "major_code",
    "password_hash",
    "app.py",
    "dataService.js",
    "run_full_pipeline_check.py",
    "HTTP 200",
    "code=0",
]


PARAGRAPH_REPLACEMENTS = {
    "（2）设计基于DIM、FACT、ADS和SYS分层结构的高校人才培养与就业数据库模型。": "（2）设计基于维度数据层、业务事实层、应用结果层和安全审计层的高校人才培养与就业数据库模型。",
    "(2) A layered database model based on DIM, FACT, ADS and SYS is designed for university talent cultivation and employment analysis.": "(2) A layered database model consisting of dimension data, business fact data, application result data and security audit data is designed for university talent cultivation and employment analysis.",
    "图2-2展示维度建模适合面向分析场景。本文以dim_student和dim_company作为维度表，以fact_academic和fact_employment作为事实表，并在此基础上建立ads系列应用结果表。该设计能够保证一个学生在学业和就业事实之间保持可追溯关联，也能够保证企业特征被就业事实复用，从而支撑专业、学校、行业、企业和地区等多维分析。": "图2-2展示维度建模适合面向分析场景。本文以学生和企业等基础对象作为维度数据，以学业过程和就业结果等业务事件作为事实数据，并在此基础上沉淀面向分析展示的应用结果数据。该设计能够保证学生在学业过程和就业结果之间保持可追溯关联，也能够保证企业特征被就业分析复用，从而支撑专业、学校、行业、企业和地区等多维分析。",
    "B/S架构将客户端统一为浏览器，适合高校、政府和公众等多角色访问[7]。本文平台前端使用React、Ant Design和ECharts构建页面，通过Axios调用Flask后端接口；后端统一读取MySQL中的基础表和ADS结果表，并返回JSON数据。": "B/S架构将客户端统一为浏览器，适合高校、政府和公众等多角色访问[7]。本文平台前端使用React、Ant Design和ECharts构建页面，通过Axios调用Flask后端服务；后端统一读取MySQL中的基础数据和应用结果数据，并以规范化形式返回页面所需信息。",
    "图2-5说明了前端调用链路。浏览器页面通过Axios携带Token访问Flask接口，后端完成角色校验和SQL查询后，将结果以code、message、data结构返回，前端再使用ECharts或Ant Design组件渲染。": "图2-5说明了前端调用链路。浏览器页面通过Axios携带Token访问Flask接口，后端完成角色校验和数据查询后，将结果封装为前端可识别的数据结构，前端再使用ECharts或Ant Design组件渲染。",
    "功能需求可概括为七类：需求预测、招生匹配、培养方案优化、就业推荐、动态监测、智能专报和统一登录审计。各功能均由数据库结果表支撑，通过Flask接口向前端提供JSON数据。": "功能需求可概括为七类：需求预测、招生匹配、培养方案优化、就业推荐、动态监测、智能专报和统一登录审计。各功能均由数据库中的分析结果支撑，并通过Flask接口向前端提供规范化数据。",
    "平台数据对象覆盖学校、专业、企业、岗位需求、毕业生、就业结果、招生计划、课程技能、政策信号、安全账号和审计日志等类型[22]。根据当前项目实现，学校与专业数据进入dim_school、dim_major_catalog和bridge_school_major；企业、行业与岗位类别进入dim_enterprise、dim_industry和dim_job_category；岗位、毕业、就业、招生、课程和政策等业务事实进入FACT层；账号、角色和审计信息进入SYS层。": "平台数据对象覆盖学校、专业、企业、岗位需求、毕业生、就业结果、招生计划、课程技能、政策信号、安全账号和审计日志等类型[22]。根据业务属性和分析需求，系统将上述数据划分为基础维度数据、业务事实数据、分析结果数据和安全审计数据四类。其中，学校、专业、行业、企业和岗位类别等相对稳定的信息作为基础维度数据，用于统一编码和跨模块关联；岗位需求、毕业生、就业结果、招生计划、课程技能和政策信号等动态信息作为业务事实数据，用于支撑预测、匹配、规则挖掘和监测分析；需求预测、招生匹配、培养优化和就业推荐等计算结果作为分析结果数据，用于前端展示和决策支持；账号、角色、访问记录和审计日志等信息作为安全审计数据，用于权限控制和系统追踪。",
    "图3-2表明，数据源层提供原始业务对象，数据治理层完成清洗、标签归一和脱敏处理，数据库层以DIM、FACT、ADS和SYS保存稳定数据，算法分析层批量写入结果表，接口服务层负责角色校验和JSON封装，应用展示层面向教师端、政府端和公众端进行差异化呈现。": "图3-2表明，数据源层提供原始业务对象，数据治理层完成清洗、标签归一和脱敏处理，数据库层按照维度数据、业务事实、分析结果和安全审计四类结构组织数据，算法分析层将预测、匹配、规则挖掘和推荐结果沉淀为可复用的分析结果，接口服务层负责身份校验、权限控制和数据封装，应用展示层面向教师端、政府端和公众端进行差异化呈现。",
    "图3-3表明，平台功能模块并非独立页面集合，而是以数据治理和ADS结果表为共同基础形成的业务闭环。需求预测结果进入招生匹配、培养优化和就业推荐模块；培养规则和专业优化建议为教师端专业建设提供解释依据；政府端和公众端在权限过滤后读取汇总结果，避免明细数据越权展示。": "图3-3表明，平台功能模块并非独立页面集合，而是以数据治理和应用结果数据为共同基础形成的业务闭环。需求预测结果进入招生匹配、培养优化和就业推荐模块；培养规则和专业优化建议为教师端专业建设提供解释依据；政府端和公众端在权限过滤后读取汇总结果，避免明细数据越权展示。",
    "数据库设计是本文平台的核心。系统采用MySQL作为关系型数据库管理系统，并借鉴数据仓库分层思想，将表结构划分为DIM维度层、FACT事实层、ADS应用结果层和SYS安全审计层[18]。维度层提供学校、专业、行业、岗位和企业等稳定描述；事实层记录招聘、毕业、就业、招生、课程和政策等业务明细；应用层保存预测、匹配、规则、优化和推荐等算法结果；安全层保存用户账号、审计日志和算法链路日志。": "数据库设计是本文平台的核心。系统采用MySQL作为关系型数据库管理系统，并借鉴数据仓库分层思想，将数据组织结构划分为维度数据层、业务事实层、应用结果层和安全审计层[18]。维度数据层提供学校、专业、行业、岗位和企业等稳定描述；业务事实层记录招聘、毕业、就业、招生、课程和政策等业务明细；应用结果层保存预测、匹配、规则、优化和推荐等算法结果；安全审计层保存用户账号、审计日志和算法运行记录。",
    "图3-4说明了数据库设计的关键原则：第一，基础业务数据与分析结果数据分离，避免模型结果反向污染原始事实；第二，学校、专业、行业和岗位等编码在各层保持一致，使跨模块追溯成为可能；第三，前端页面优先读取ADS应用层，保证展示性能和接口稳定性；第四，安全审计表独立于业务表，便于后续扩展账号策略、操作追踪和风险审计功能。": "图3-4说明了数据库设计的关键原则：第一，基础业务数据与分析结果数据分离，避免模型结果反向污染原始事实；第二，学校、专业、行业和岗位等编码在各层保持一致，使跨模块追溯成为可能；第三，前端页面优先读取应用结果数据，保证展示性能和接口稳定性；第四，安全审计数据独立于业务数据，便于后续扩展账号策略、操作追踪和风险审计功能。",
    "维度表采用主键编码和常用查询字段索引。例如dim_school以school_id作为主键，dim_major_catalog以major_code作为主键，bridge_school_major以school_id和major_code组成联合主键，用于描述某高校是否开设某专业、是否为优势专业以及历史招生规模。行业与岗位维度通过industry_id和job_category_id形成层级关系，企业维度记录企业规模、所有制、高新技术、专精特新和上海重点产业标签，为就业质量分析和推荐计算提供画像。": "维度数据采用唯一编码和常用查询条件索引。系统为学校、专业和学校—专业关联关系设置稳定标识，用于描述某高校是否开设某专业、是否为优势专业以及历史招生规模。行业与岗位类别之间保持层级关系，企业信息记录企业规模、所有制、高新技术、专精特新和上海重点产业标签，为就业质量分析和推荐计算提供画像。",
    "事实表围绕业务事件展开。fact_job_posting按月份记录企业岗位需求人数、薪资区间、学历要求、技能标签和偏好专业；fact_graduate记录毕业生学校、专业、学历、技能、实习和求职意向；fact_employment记录毕业生就业单位、行业、岗位、薪资、专业相关度和社保核验状态；fact_enrollment_plan、fact_course_skill和fact_policy_signal分别支撑招生分析、培养规则挖掘和政策热度计算。通过这些事实表，系统能够从需求侧、供给侧和政策侧同时刻画高校人才培养环境。": "业务事实数据围绕具体业务事件展开。岗位需求数据按月份记录企业需求人数、薪资区间、学历要求、技能标签和偏好专业；毕业生信息记录学校、专业、学历、技能、实习和求职意向；就业结果数据记录就业单位、行业、岗位、薪资、专业相关度和社保核验状态；招生计划、课程技能和政策信号数据分别支撑招生分析、培养规则挖掘和政策热度计算。通过这些业务事实数据，系统能够从需求侧、供给侧和政策侧同时刻画高校人才培养环境。",
    "应用结果表采用“按场景冗余关键展示字段”的方式设计。例如，ads_job_demand_forecast不仅保存school_id、major_code等编码，还保存school_name、major_name、industry_name和job_category_name，便于前端直接展示；ads_job_recommendation保存推荐岗位、企业、相似度、置信等级和推荐理由，避免页面再进行复杂拼接；ads_major_optimization保存建议类型、优先级、证据分、原因说明和相关指标，使教师端能够解释专业扩招、稳招、缩招或加强实践培养等建议的依据。": "应用结果数据采用“按场景冗余关键展示信息”的方式设计。例如，岗位需求预测结果同时保留学校、专业、行业和岗位类别等核心标识及其展示名称，便于前端直接呈现；就业推荐结果保存推荐岗位、企业、匹配程度、置信等级和推荐理由，避免页面再进行复杂拼接；专业优化结果保存建议类型、优先级、证据分、原因说明和相关指标，使教师端能够解释专业扩招、稳招、缩招或加强实践培养等建议的依据。",
    "在关键表字段设计上，本文优先列出支撑核心业务链路的代表字段。字段选择以当前scripts/create_tables.py中的DDL为依据，既覆盖业务主键，也覆盖接口查询和页面展示所需字段。": "在核心数据对象设计上，本文优先列出支撑核心业务链路的主要属性。属性选择围绕业务识别、分析计算和页面展示需要展开，既覆盖业务事件的唯一标识，也覆盖跨模块查询和结果解释所需信息。",
    "表3-6说明，平台数据库以school_id、major_code、industry_id、job_category_id和month等字段贯通业务事实和算法结果。对于多学校、多专业和多岗位组合，统一编码可以显著降低接口聚合和前端展示的复杂度。": "表3-6说明，平台数据库通过学校、专业、行业、岗位类别和时间等核心标识贯通业务事实与算法结果。对于多学校、多专业和多岗位组合的综合分析场景，统一编码和统一时间粒度能够降低跨表关联和结果展示的复杂度，提高数据追溯和模块复用能力。",
    "在约束与索引设计方面，dim_school、dim_major_catalog、dim_industry、dim_job_category和dim_enterprise均设置主键；bridge_school_major采用school_id与major_code联合主键；事实表使用posting_id、graduate_id、employment_id等业务主键。常用过滤字段包括school_id、major_code、industry_id、job_category_id、month、forecast_month和user_id，并在DDL中设置相应索引。当前实现主要通过主键和索引保证查询效率，外键关系以逻辑约束和脚本校验为主，便于原型系统快速重建和批量导入。": "在约束与索引设计方面，系统为学校、专业、行业、岗位类别和企业等基础对象设置唯一标识，并通过学校与专业之间的关联关系描述高校专业开设情况。业务事实数据以岗位需求、毕业生、就业结果等业务事件为基本记录单元，应用结果数据则以分析场景和展示需求为组织依据。对于学校、专业、行业、岗位类别和时间等高频查询条件，系统设置相应索引，以提高跨模块查询和页面展示效率。当前实现主要通过唯一标识、逻辑约束和数据校验保证数据一致性，便于原型系统快速重建和批量导入。",
    "项目目录也体现了上述分工。scripts目录保存数据生成、建表、导入、算法计算和验收脚本；app.py提供主要Flask接口；backend目录保存安全、提示词和大模型调用相关代码；frontend目录保存React前端页面、组件、路由、数据服务和图表适配逻辑；data/generated目录保存生成后的CSV和验收报告；docs目录保存接口、页面和全链路检查报告。": "从工程结构看，系统按照数据处理、算法分析、后端服务、前端展示和文档报告等功能进行模块化组织。数据处理模块负责数据生成、数据库初始化和批量导入；算法分析模块负责特征加工、需求预测、招生匹配、培养优化和就业推荐；后端服务模块负责接口封装、权限校验和数据返回；前端展示模块负责页面路由、图表展示和交互逻辑；文档报告模块用于保存接口说明、测试结果和系统运行记录。",
    "前端通过统一的dataService封装Axios请求，并在请求拦截器中自动携带Token。后端接口统一返回code、message和data结构，前端通过数据适配器将不同接口结果转换为图表组件所需格式。该设计降低了页面与数据库字段之间的直接耦合，也便于后续在不改动页面的情况下增加接口缓存、异常提示和空数据状态处理。": "前端通过统一的数据服务层封装Axios请求，并在请求拦截器中自动携带身份凭证。后端接口对结果进行统一封装，前端通过数据适配器将不同业务结果转换为图表组件所需格式。该设计降低了页面与数据库属性之间的直接耦合，也便于后续在不改动页面的情况下增加接口缓存、异常提示和空数据状态处理。",
    "由于真实高校学生和就业数据具有较强隐私属性，本文在系统实现阶段使用脱敏仿真数据验证平台链路。数据生成脚本platform_data_factory.py依据上海高校、专业目录、行业岗位分类、企业画像、政策方向和季节性招聘规律构造样本数据。生成内容覆盖学校、专业、行业、企业、岗位需求、毕业生、就业结果、招生计划、课程技能和政策信号等对象，能够支撑后续预测、匹配、规则挖掘和推荐计算[13]。": "由于真实高校学生和就业数据具有较强隐私属性，本文在系统实现阶段使用脱敏仿真数据验证平台链路。数据生成模块依据上海高校、专业目录、行业岗位分类、企业画像、政策方向和季节性招聘规律构造样本数据。生成内容覆盖学校、专业、行业、企业、岗位需求、毕业生、就业结果、招生计划、课程技能和政策信号等对象，能够支撑后续预测、匹配、规则挖掘和推荐计算[13]。",
    "如图4-1所示，数据处理链路首先生成标准CSV文件，然后通过create_tables.py创建数据库表结构和索引，再由PutData.py批量导入MySQL。导入后，质量检查脚本会验证学校数量、专业目录覆盖、岗位行业分布、就业率差异、薪资差异、岗位需求季节波动和模型结果非负等条件。质量检查通过后，Spark-all.py对事实表进行多表关联和聚合，生成ads_job_demand_features等应用层输入。": "如图4-1所示，数据处理链路首先生成标准化数据文件，然后完成数据库结构初始化和批量数据导入。数据导入后，系统从学校数量、专业覆盖、岗位行业分布、就业率差异、薪资差异、岗位需求季节波动和模型结果合理性等方面进行质量检查。质量检查通过后，特征聚合模块对业务事实数据进行多表关联和汇总，形成后续需求预测、招生匹配和培养优化所需的基础特征。",
    "本次全链路验收结果显示，基础数据和结果数据均已成功生成并写入数据库。其中岗位需求事实表达到81000条，毕业生事实表达到51000条，就业事实表达到42000条，Spark聚合后生成127501条岗位需求特征记录。质量检查中，计算机相关岗位占比为0.1075，非计算机岗位覆盖18个行业，岗位需求曲线季节波动系数为2.063，说明样本具有一定业务差异性和趋势波动特征。": "本次全链路验收结果显示，基础数据和结果数据均已成功生成并写入数据库。其中岗位需求、毕业生和就业结果等业务记录分别达到81000条、51000条和42000条，特征聚合后形成127501条岗位需求特征记录。质量检查中，计算机相关岗位占比为0.1075，非计算机岗位覆盖18个行业，岗位需求曲线季节波动系数为2.063，说明样本具有一定业务差异性和趋势波动特征。",
    "平台核心算法链路由特征聚合、需求预测、招生匹配、培养规则挖掘、就业推荐和动态监测六个阶段组成。各阶段均以数据库表为输入和输出，不依赖前端页面实时计算。算法脚本运行后会将状态、输入表、输出表、算法名称、记录数和耗时写入ads_algorithm_chain_log，便于在算法链路页面中展示运行情况。": "平台核心算法链路由特征聚合、需求预测、招生匹配、培养规则挖掘、就业推荐和动态监测六个阶段组成。各阶段均以数据库中的业务数据和分析结果为输入与输出，不依赖前端页面实时计算。算法运行后会记录运行状态、输入输出范围、算法名称、记录数量和耗时等信息，便于在算法链路页面中展示运行情况。",
    "Spark特征聚合模块读取fact_job_posting、fact_employment、fact_enrollment_plan和fact_policy_signal等表，对岗位需求、就业质量、招生热度和政策热度进行统一加工。聚合粒度为月份、学校、专业、行业和岗位类别，输出字段包括demand_count_sum、posting_count、avg_salary、policy_heat、employment_rate、match_score、skill_gap_score、enrollment_pressure和school_major_strength_score等。这些字段既服务于需求预测，也为招生匹配和培养优化提供基础特征。": "Spark特征聚合模块以岗位需求、就业结果、招生计划和政策信号等业务事实数据为输入，对岗位需求强度、就业质量、招生热度、政策热度和学校专业基础等指标进行统一加工。聚合粒度设置为月份、学校、专业、行业和岗位类别，使系统能够同时刻画时间变化、学校差异、专业差异和行业需求差异。聚合后的特征结果既服务于岗位需求预测，也为招生匹配和培养优化提供基础数据支撑。",
    "需求预测模块读取ads_job_demand_features，选择样本月份较完整的学校—专业—行业—岗位组合，构造未来12个月的预测结果。工程实现中采用LSTM时序预测思想，并结合季节性回退策略处理样本不足和局部波动问题。模型不仅输出predicted_demand_count，还给出lower_bound、upper_bound、demand_growth_rate、demand_level、mape和track_rank等字段，便于页面展示预测区间、需求等级和模型解释。": "需求预测模块在特征聚合结果的基础上，选择历史月份较完整的学校—专业—行业—岗位组合，构造未来12个月的岗位需求预测结果。工程实现中借鉴LSTM时间序列预测思想，并结合季节性回退策略处理样本不足和局部波动问题。模型结果不仅包括未来需求规模，还包括预测区间、增长水平、需求等级和误差评价等指标，便于页面展示趋势变化并为管理者提供可解释的判断依据。",
    "招生匹配模块将专业视为分析对象，综合未来需求、就业率、平均薪资、招生热度、政策热度和学校专业强度计算match_score[14]。对于每个学校专业，系统输出建议类型、建议动作和原因说明。例如当需求增长和就业质量较高时，专业可能被建议扩招；当就业率和需求增长均不足时，可能被建议缩招；当政策热度高且学校专业基础较好时，可能被建议重点扶持。": "招生匹配模块将专业视为分析对象，综合未来需求、就业率、平均薪资、招生热度、政策热度和学校专业基础计算匹配程度[14]。对于每个学校专业，系统输出建议类型、建议动作和原因说明。例如当需求增长和就业质量较高时，专业可能被建议扩招；当就业率和需求增长均不足时，可能被建议缩招；当政策热度高且学校专业基础较好时，可能被建议重点扶持。",
    "培养优化模块由两部分组成。首先，FPgrowth-all.py从课程技能、就业结果和岗位类别中挖掘“课程—技能—岗位”关联规则，输出support、confidence、lift和evidence_score等指标；其次，CF-all.py将需求预测、就业质量、招生热度、政策热度和规则证据综合为专业结构优化建议，写入ads_major_optimization。该实现方式使培养建议能够回溯到规则证据和业务指标，避免只给出结论而缺乏解释。": "培养优化模块由关联规则挖掘和专业结构优化两部分组成。首先，系统从课程技能、就业结果和岗位类别等数据中挖掘“课程—技能—岗位”之间的潜在关联，并通过支持度、置信度、提升度和证据分等指标衡量规则可靠性。其次，系统综合需求预测、就业质量、招生热度、政策热度和规则证据，形成面向专业结构调整和培养方案优化的建议。该实现方式使培养建议能够回溯到规则证据和业务指标，避免只给出结论而缺乏解释。",
    "就业推荐模块读取毕业生画像、岗位预测结果、岗位技能标签、企业属性和历史就业质量数据，将学生技能、求职意向、岗位需求强度、专业基础和企业吸引力组合为相似度分数[30]。每名学生输出Top-K推荐岗位和企业，并给出similarity_score、confidence_level、salary_reference和recommendation_reason等字段。推荐理由采用结构化文本，说明学生专业与岗位技能、求职意向和未来需求之间的匹配关系。": "就业推荐模块读取毕业生画像、岗位预测结果、岗位技能标签、企业属性和历史就业质量数据，将学生技能、求职意向、岗位需求强度、专业基础和企业吸引力组合为匹配程度[30]。系统为每名学生生成若干个候选岗位推荐结果，并给出匹配程度、置信等级、薪资参考和推荐理由等信息。推荐理由采用结构化文本，说明学生专业与岗位技能、求职意向和未来需求之间的匹配关系。",
    "动态监测模块主要面向政府端和公众端。政府端可以查看跨校就业率、平均薪资、高质量就业比例、重点产业就业占比、学校对标和区域风险预警；公众端只展示脱敏公开汇总信息，如院校对比、专业趋势和就业方向建议。由于监测页面直接读取ADS或汇总接口结果，因此既保证了响应速度，也避免了公众端接触敏感明细数据。": "动态监测模块主要面向政府端和公众端。政府端可以查看跨校就业率、平均薪资、高质量就业比例、重点产业就业占比、学校对标和区域风险预警；公众端只展示脱敏公开汇总信息，如院校对比、专业趋势和就业方向建议。由于监测页面直接读取应用结果数据或汇总结果，因此既保证了响应速度，也避免了公众端接触敏感明细数据。",
    "后端接口集中由根目录app.py提供，backend/app.py作为启动入口导入app对象并执行启动检查。接口统一使用/api路径前缀，返回结构为code、message和data。前端dataService.js通过Axios封装调用，并在请求拦截器中携带Token。": "后端接口采用统一的服务入口进行管理，主要负责数据查询、身份认证、权限校验、结果封装和异常处理。系统接口按照业务模块进行组织，包括需求预测、招生匹配、培养优化、就业推荐、动态监测、智能专报和安全审计等类型。前端通过统一的数据服务层调用后端接口，并在用户登录后自动携带身份凭证，从而实现页面访问与角色权限之间的关联。",
    "表4-4中的接口名称均来自当前app.py和dataService.js。API合约检查报告显示，教师端、政府端和公众端主要接口均能返回HTTP 200和code=0，且核心数据不存在全0风险。": "表4-4中的接口类别按照系统业务模块进行整理。接口测试结果表明，教师端、政府端和公众端的核心接口能够正常返回数据，主要业务页面能够完成数据加载和结果展示，说明后端服务与前端页面之间的数据链路具备基本可用性。",
    "平台涉及学生就业、薪资、企业和推荐结果等敏感数据，当前实现通过账号表、Token、角色权限和审计日志进行基础保护[40]。sys_user_account保存password_hash、hash_algo、role和school_id等字段；登录成功后后端签发Token；前端请求时自动携带Token；后端根据角色限制教师端、政府端和公众端的数据范围。": "平台涉及学生就业、薪资、企业信息和推荐结果等敏感数据，当前实现通过账号认证、身份令牌、角色权限和审计日志进行基础保护[40]。系统在用户登录后生成身份凭证，前端访问受保护接口时自动携带该凭证，后端根据用户角色和所属范围限制可访问数据。教师端主要访问本校或本专业相关数据，政府端访问跨学校汇总数据，公众端仅访问脱敏后的公开信息。通过上述机制，系统在原型阶段实现了基本的身份认证、权限隔离和访问留痕。",
    "全链路验收脚本run_full_pipeline_check.py共执行15个步骤，覆盖数据生成、建表、导入、Spark聚合、岗位需求预测、招生匹配、培养方案优化、就业推荐、安全初始化、数据质量检查、页面逻辑检查、API合约检查、前端静态扫描、依赖安装和前端构建。验收报告显示总体结论为通过，阻断性问题为0。": "为验证系统整体可运行性，本文对数据生成、数据库初始化、数据导入、特征聚合、需求预测、招生匹配、培养方案优化、就业推荐、安全初始化、数据质量检查、接口连通性检查和前端构建等环节进行了全链路测试。测试结果表明，系统各主要模块能够按照预期顺序完成运行，核心结果能够写入数据库并被前端页面读取，未发现影响系统基本运行的阻断性问题。",
    "本章从开发环境、数据入库、算法链路、接口服务、前端页面和安全审计等方面说明了一体化平台的实现过程。当前系统以MySQL分层数据库为核心，以脚本化算法计算形成ADS结果表，以Flask接口和React页面完成多角色展示，形成了从数据治理到业务可视化的闭环实现。": "本章从开发环境、数据入库、算法链路、接口服务、前端页面和安全审计等方面说明了一体化平台的实现过程。当前系统以MySQL分层数据库为核心，通过脚本化算法计算形成应用结果数据，并以Flask接口和React页面完成多角色展示，形成了从数据治理到业务可视化的闭环实现。",
    "从系统实现角度看，本文采用B/S架构进行平台开发。后端基于Flask提供统一接口服务，承担数据查询、权限认证、审计留痕、结果封装和专报生成等功能；前端基于React、Ant Design和ECharts构建教师端、政府端和公众端等差异化页面，实现需求趋势、招生匹配、培养规则、就业推荐、学校对标和公开展示等功能；数据库层借鉴数据仓库思想，设计DIM维度层、FACT事实层、ADS应用结果层和SYS安全审计层，实现基础业务数据、算法分析结果和系统安全数据的统一管理。系统通过脱敏仿真数据完成全链路验证，验证内容覆盖数据生成、建表入库、特征加工、模型计算、接口调用、页面构建和安全审计等环节，结果表明平台具备一定的可运行性、可解释性和可扩展性。": "从系统实现角度看，本文采用B/S架构进行平台开发。后端基于Flask提供统一接口服务，承担数据查询、权限认证、审计留痕、结果封装和专报生成等功能；前端基于React、Ant Design和ECharts构建教师端、政府端和公众端等差异化页面，实现需求趋势、招生匹配、培养规则、就业推荐、学校对标和公开展示等功能；数据库层借鉴数据仓库思想，设计维度数据层、业务事实层、应用结果层和安全审计层，实现基础业务数据、算法分析结果和系统安全数据的统一管理。系统通过脱敏仿真数据完成全链路验证，验证内容覆盖数据生成、建表入库、特征加工、模型计算、接口调用、页面构建和安全审计等环节，结果表明平台具备一定的可运行性、可解释性和可扩展性。",
    "本文以关系型数据库为核心底座，借鉴数据仓库分层思想，将数据库划分为DIM维度层、FACT事实层、ADS应用结果层和SYS安全审计层。DIM层主要保存学校、专业、行业、岗位类别和企业等稳定维度信息；FACT层保存岗位需求、毕业生、就业、招生、课程和政策等业务事实；ADS层保存需求预测、招生匹配、培养规则、专业优化和就业推荐等算法结果；SYS层保存用户账号、角色权限、审计日志和算法链路日志。该分层设计实现了基础数据与分析结果的分离，避免算法结果直接影响原始业务事实，同时通过统一编码实现不同模块之间的数据共享和结果追溯。前端页面和后端接口优先读取ADS应用结果表，既提高了查询效率，也使算法计算结果能够沉淀为可查询、可解释、可复用的数据资产。": "本文以关系型数据库为核心底座，借鉴数据仓库分层思想，将数据库划分为维度数据层、业务事实层、应用结果层和安全审计层。维度数据层主要保存学校、专业、行业、岗位类别和企业等稳定维度信息；业务事实层保存岗位需求、毕业生、就业、招生、课程和政策等业务事实；应用结果层保存需求预测、招生匹配、培养规则、专业优化和就业推荐等算法结果；安全审计层保存用户账号、角色权限、审计日志和算法运行记录。该分层设计实现了基础数据与分析结果的分离，避免算法结果直接影响原始业务事实，同时通过统一编码实现不同模块之间的数据共享和结果追溯。前端页面和后端接口优先读取应用结果数据，既提高了查询效率，也使算法计算结果能够沉淀为可查询、可解释、可复用的数据资产。",
    "本文基于脱敏仿真数据对平台进行了全链路测试与验证。测试内容包括数据生成、数据库建表、数据导入、特征聚合、需求预测、招生匹配、培养优化、就业推荐、安全初始化、质量检查、接口合约检查、前端静态扫描和项目构建等环节。实验结果表明，系统能够完成十万级样本数据的组织与分析，核心算法结果能够写入ADS结果表，主要接口能够正常返回数据，前端页面能够完成主要业务图表和结果列表展示。通过审计日志和算法链路日志，系统还能够记录关键访问行为和算法运行状态，为后续系统维护、异常排查和结果复核提供依据。": "本文基于脱敏仿真数据对平台进行了全链路测试与验证。测试内容包括数据生成、数据库初始化、数据导入、特征聚合、需求预测、招生匹配、培养优化、就业推荐、安全初始化、质量检查、接口连通性检查和项目构建等环节。实验结果表明，系统能够完成十万级样本数据的组织与分析，核心算法结果能够写入应用结果数据，主要接口能够正常返回数据，前端页面能够完成主要业务图表和结果列表展示。通过审计日志和算法运行记录，系统还能够记录关键访问行为和算法运行状态，为后续系统维护、异常排查和结果复核提供依据。",
    "高校人才培养与就业治理涉及时间趋势、结构匹配、能力关联和个性化推荐等多类问题，单一算法难以覆盖全部业务场景。针对这一特点，本文将不同算法与不同治理任务相结合。需求预测场景具有明显的时间连续性和季节波动特征，本文借鉴LSTM时间序列预测思想，用于刻画人才需求随时间变化的趋势规律；培养优化场景需要揭示课程体系、技能标签和就业结果之间的隐性关联，本文采用关联规则挖掘方法，分析“课程—技能—岗位—就业质量”之间的共现关系；就业推荐场景强调学生画像与岗位画像之间的匹配程度，本文结合协同过滤思想和相似度计算方法，为学生提供个性化岗位推荐和推荐理由。不同算法结果最终沉淀到ADS应用结果表中，并通过统一接口提供给前端页面和智能专报模块使用，从而增强了平台分析结果的可解释性和可复用性。": "高校人才培养与就业治理涉及时间趋势、结构匹配、能力关联和个性化推荐等多类问题，单一算法难以覆盖全部业务场景。针对这一特点，本文将不同算法与不同治理任务相结合。需求预测场景具有明显的时间连续性和季节波动特征，本文借鉴LSTM时间序列预测思想，用于刻画人才需求随时间变化的趋势规律；培养优化场景需要揭示课程体系、技能标签和就业结果之间的隐性关联，本文采用关联规则挖掘方法，分析“课程—技能—岗位—就业质量”之间的共现关系；就业推荐场景强调学生画像与岗位画像之间的匹配程度，本文结合协同过滤思想和相似度计算方法，为学生提供个性化岗位推荐和推荐理由。不同算法结果最终沉淀到应用结果数据中，并通过统一接口提供给前端页面和智能专报模块使用，从而增强了平台分析结果的可解释性和可复用性。",
    "在技术实现层面，本文围绕高校大数据治理和动态监测需求，构建了涵盖数据存储、特征加工、算法计算、接口服务和前端展示的一体化技术栈。平台采用分层数据组织方式承载基础数据和分析结果，并结合Spark并行计算思想完成多表关联、特征聚合和结果加工；后端通过Flask提供轻量级接口服务，实现数据查询、权限认证和结果封装；前端通过React和ECharts实现多角色可视化展示；安全层通过账号权限、Token认证和审计日志记录用户访问行为。该技术路线兼顾了原型系统的可实现性、功能模块的可维护性和未来扩展到分布式环境的可迁移性。通过脚本化数据链路和ADS结果表设计，平台将算法计算过程与在线页面访问过程解耦，降低了在线接口的计算压力，提高了系统的工程适配性。": "在技术实现层面，本文围绕高校大数据治理和动态监测需求，构建了涵盖数据存储、特征加工、算法计算、接口服务和前端展示的一体化技术栈。平台采用分层数据组织方式承载基础数据和分析结果，并结合Spark并行计算思想完成多表关联、特征聚合和结果加工；后端通过Flask提供轻量级接口服务，实现数据查询、权限认证和结果封装；前端通过React和ECharts实现多角色可视化展示；安全层通过账号权限、Token认证和审计日志记录用户访问行为。该技术路线兼顾了原型系统的可实现性、功能模块的可维护性和未来扩展到分布式环境的可迁移性。通过脚本化数据链路和应用结果数据设计，平台将算法计算过程与在线页面访问过程解耦，降低了在线接口的计算压力，提高了系统的工程适配性。",
}

CELL_REPLACEMENTS = {
    "ADS结果表、索引、离线聚合、接口缓存预留": "应用结果数据、索引、离线聚合、接口缓存预留",
    "密码哈希、角色字段、失败限制、审计日志": "密码加密、角色权限、失败限制、审计日志",
    "编码字典、字段映射、主键约束、结果表复用": "统一编码、属性映射、唯一约束、结果复用",
    "分层数据库、模块化接口、脚本化数据链路": "分层数据库、模块化接口、自动化数据链路",
    "代表表": "代表数据对象",
    "维度表与事实表": "维度数据与业务事实数据",
    "join、groupBy、窗口统计": "关联处理、分组聚合和窗口统计",
    "模型训练特征表": "模型训练特征数据",
    "学校/专业/行业维度聚合": "按学校、专业和行业等维度聚合",
    "ads_employment_summary": "就业汇总分析结果",
    "dim_school、dim_major_catalog、bridge_school_major": "学校基础信息、专业目录数据、学校专业开设关系",
    "dim_enterprise、dim_industry、dim_job_category、fact_job_posting": "企业基础信息、行业基础信息、岗位类别数据、岗位需求数据",
    "fact_graduate、fact_employment": "毕业生信息、就业结果数据",
    "fact_enrollment_plan、fact_course_skill": "招生计划数据、课程技能数据",
    "fact_enrollment_plan、fact_course_skill、fact_policy_signal": "招生计划数据、课程技能数据、政策信号数据",
    "fact_policy_signal、sys_user_account、sys_audit_log": "政策信号数据、用户账号数据、审计日志数据",
    "DIM、FACT、SYS基础表及清洗日志": "维度数据、业务事实数据、安全审计数据及清洗记录",
    "ads_job_demand_features、ads_job_demand_forecast": "需求预测特征数据和岗位需求预测结果",
    "ads_enrollment_matching及匹配解释": "招生匹配结果及匹配解释",
    "ads_training_rules、ads_major_optimization": "培养规则结果和专业优化结果",
    "ads_job_recommendation及推荐理由": "就业推荐结果及推荐理由",
    "ADS结果表、学校对标表、接口聚合结果": "应用结果数据、学校对标结果、接口聚合结果",
    "Token、审计日志、访问记录": "身份令牌、审计日志、访问记录",
    "DIM维度层": "维度数据层",
    "FACT事实层": "业务事实层",
    "ADS应用层": "应用结果层",
    "SYS安全层": "安全审计层",
    "dim_school、dim_major_catalog、bridge_school_major、dim_industry、dim_job_category、dim_enterprise": "学校信息、专业信息、学校专业开设关系、行业信息、岗位类别信息、企业信息",
    "fact_job_posting、fact_graduate、fact_employment、fact_enrollment_plan、fact_course_skill、fact_policy_signal": "岗位需求数据、毕业生信息、就业结果数据、招生计划数据、课程技能数据、政策信号数据",
    "ads_job_demand_features、ads_job_demand_forecast、ads_enrollment_matching、ads_training_rules、ads_major_optimization、ads_job_recommendation": "需求预测特征、岗位需求预测结果、招生匹配结果、培养规则结果、专业优化结果、就业推荐结果",
    "sys_user_account、sys_audit_log、ads_algorithm_chain_log": "用户账号数据、审计日志数据、算法运行记录",
    "主要数据表": "主要数据对象",
    "表名": "数据对象",
    "关键字段": "主要属性",
    "设计说明": "作用说明",
    "dim_school": "学校信息",
    "school_id、school_name、city、district、major_count": "学校名称、学校类型、所属区域、办学层次、专业数量",
    "学校维度主表，支撑学校范围过滤和跨校对比": "支撑学校对标、区域分析和权限过滤",
    "dim_major_catalog": "专业信息",
    "major_code、major_name、discipline_category、major_class": "专业名称、专业类别、所属学科、专业状态",
    "专业目录维表，统一专业编码和学科门类": "支撑专业分析、招生匹配和培养优化",
    "bridge_school_major": "学校—专业关联数据",
    "school_id、major_code、school_major_strength_score": "开设状态、优势标识、历史招生规模、专业基础",
    "校专桥接表，表示高校开设专业及专业强度": "描述高校专业开设关系，支撑专业基础评价",
    "fact_job_posting": "岗位需求",
    "posting_id、enterprise_id、industry_id、job_category_id、month、demand_count": "岗位类别、需求人数、薪资水平、技能要求、发布时间",
    "岗位需求事实表，支撑月度需求聚合": "支撑需求预测和市场趋势分析",
    "fact_graduate": "毕业生信息",
    "graduate_id、school_id、major_code、skill_tags、job_intention_tags": "所属学校、所属专业、能力标签、求职意向",
    "毕业生画像事实表，支撑就业推荐": "支撑学生画像和就业推荐",
    "fact_employment": "就业结果",
    "employment_id、graduate_id、school_id、major_code、salary、match_score": "就业单位、岗位类别、薪资水平、专业相关度、就业质量",
    "就业结果事实表，支撑就业质量与规则分析": "支撑就业质量评价和培养反馈",
    "fact_enrollment_plan": "招生计划",
    "plan_id、school_id、major_code、year、planned_quota、applicant_count": "招生年份、专业方向、计划人数、生源特征",
    "招生计划事实表，支撑招生匹配": "支撑招生匹配和专业结构调整",
    "fact_course_skill": "课程技能",
    "course_id、school_id、major_code、course_name、skill_tags": "课程模块、技能标签、能力要求、实践环节",
    "课程技能事实表，支撑培养规则挖掘": "支撑培养方案优化和规则挖掘",
    "ads_job_demand_forecast": "需求预测结果",
    "forecast_month、predicted_demand_count、mape、track_rank": "预测月份、预测需求、增长趋势、需求等级",
    "需求预测结果表，支撑需求趋势页面": "支撑岗位需求趋势展示和专业调整",
    "ads_enrollment_matching": "招生匹配结果",
    "match_score、recommendation_action、precision_at_k": "匹配程度、建议动作、评价指标",
    "招生匹配结果表，支撑专业匹配建议": "支撑招生规模判断和专业匹配建议",
    "ads_major_optimization": "专业优化结果",
    "primary_suggestion_type、priority_score、suggestion_reason": "建议类型、优先级、证据分、原因说明",
    "专业优化结果表，支撑培养优化页面": "支撑专业结构调整和培养优化展示",
    "ads_job_recommendation": "就业推荐结果",
    "graduate_id、enterprise_id、similarity_score、confidence_level": "推荐岗位、匹配程度、置信等级、推荐理由",
    "就业推荐结果表，支撑Top-K推荐": "支撑个性化就业指导",
    "sys_user_account": "安全审计信息",
    "user_id、username、password_hash、role、school_id": "用户角色、所属范围、访问凭证、权限状态",
    "用户账号表，支撑登录和角色权限": "支撑身份认证、权限控制和访问追踪",
    "sys_audit_log": "审计日志信息",
    "audit_id、user_id、action、module、request_path、created_at": "访问时间、操作类型、访问模块、运行状态",
    "审计日志表，支撑访问追踪": "支撑日志追踪、系统审计和异常复核",
    "MySQL、维度表、事实表、ADS结果表": "MySQL、维度数据、业务事实数据、应用结果数据",
    "Token、密码哈希、角色权限、sys_audit_log": "身份令牌、密码加密、角色权限、审计日志",
    "代表数据表": "代表数据对象",
    "dim_industry、dim_job_category、dim_enterprise": "行业基础信息、岗位类别数据、企业基础信息",
    "fact_job_posting、fact_graduate、fact_employment": "岗位需求数据、毕业生信息、就业结果数据",
    "ads_job_demand_features、ads_job_demand_forecast、ads_job_recommendation": "需求预测特征数据、岗位需求预测结果、就业推荐结果",
    "输出表或指标": "输出结果或指标",
    "ads_job_demand_forecast、MAPE": "岗位需求预测结果、误差评价",
    "ads_enrollment_matching、Precision@K": "招生匹配结果、Precision@K",
    "ads_training_rules、evidence_score": "培养规则结果、证据分",
    "ads_major_optimization": "专业优化结果",
    "ads_job_recommendation、Top1相似度": "就业推荐结果、Top1相似度",
    "接口名称": "接口内容",
    "/api/auth/login、/api/auth/me、/api/auth/logout": "登录、会话校验和退出功能",
    "sys_user_account、sys_audit_log": "用户账号数据和审计日志数据",
    "/api/demand/kpi、/api/demand/forecast、/api/demand/forecast/eval、/api/demand/forecast/backtest": "需求概览、需求预测、模型评价和历史回测功能",
    "ads_job_demand_forecast及评估表": "岗位需求预测结果及评价数据",
    "/api/enrollment/matching、/api/enrollment-matching-evaluation": "招生匹配和匹配效果评价功能",
    "ads_enrollment_matching": "招生匹配结果",
    "/api/major/optimization、/api/training/rules、/api/training-program-optimization": "专业优化、培养规则和培养方案优化功能",
    "ads_major_optimization、ads_training_rules": "专业优化结果和培养规则结果",
    "/api/recommendation/summary、/api/recommendation/jobs、/api/recommendation/student": "推荐概览、岗位推荐和学生推荐功能",
    "ads_job_recommendation": "就业推荐结果",
    "/api/employment-summary、/api/monitor/school、/api/public/salary-ranking": "就业汇总、学校监测和公开薪资排行功能",
    "/api/report/ai、/api/algorithm/chain-log、/api/audit/logs": "智能专报、算法运行记录和审计日志功能",
    "ADS结果表、ads_algorithm_chain_log、sys_audit_log": "应用结果数据、算法运行记录和审计日志数据",
}


def find_source() -> Path:
    named = Path("毕设.docx")
    if named.exists():
        return named
    candidates = [p for p in Path(".").glob("*.docx") if p.stat().st_size == SOURCE_SIZE]
    if candidates:
        return candidates[0]
    raise FileNotFoundError("未找到 毕设.docx 或已知格式统一版本源文件")


def term_pattern(terms: list[str]) -> re.Pattern[str]:
    return re.compile("|".join(re.escape(t) for t in sorted(terms, key=len, reverse=True)))


def count_hits(doc: Document, pattern: re.Pattern[str]) -> Counter:
    counts: Counter = Counter()
    for para in doc.paragraphs:
        for hit in pattern.findall(para.text):
            counts[hit] += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for hit in pattern.findall(para.text):
                        counts[hit] += 1
    return counts


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        first = paragraph.runs[0]
        for run in paragraph.runs[1:]:
            paragraph._element.remove(run._element)
        first.text = text
    else:
        paragraph.add_run(text)


def replace_exact_paragraphs(doc: Document) -> int:
    changed = 0
    for para in doc.paragraphs:
        text = para.text
        if text in PARAGRAPH_REPLACEMENTS:
            set_paragraph_text(para, PARAGRAPH_REPLACEMENTS[text])
            changed += 1
            continue
        new = text
        for old, repl in PARAGRAPH_REPLACEMENTS.items():
            if old in new:
                new = new.replace(old, repl)
        if new != text:
            set_paragraph_text(para, new)
            changed += 1
    return changed


def replace_cell_text(cell, new_text: str) -> None:
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.add_paragraph(new_text)
        return
    set_paragraph_text(paragraphs[0], new_text)
    for para in paragraphs[1:]:
        set_paragraph_text(para, "")


def replace_table_cells(doc: Document) -> int:
    changed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                new = text
                if text in CELL_REPLACEMENTS:
                    new = CELL_REPLACEMENTS[text]
                else:
                    for old, repl in CELL_REPLACEMENTS.items():
                        if old in new:
                            new = new.replace(old, repl)
                if new != text:
                    replace_cell_text(cell, new)
                    changed += 1
    return changed


def residual_locations(doc: Document) -> list[str]:
    locations: list[str] = []
    for term in RESIDUAL_TERMS:
        lower_term = term.lower()
        found = False
        for i, para in enumerate(doc.paragraphs):
            if lower_term in para.text.lower():
                locations.append(f"{term}: 正文段落{i + 1}")
                found = True
                break
        if found:
            continue
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    if lower_term in cell.text.lower():
                        locations.append(f"{term}: 表格{ti + 1}第{ri + 1}行第{ci + 1}列")
                        found = True
                        break
                if found:
                    break
            if found:
                break
    return locations


def count_figures_tables(doc: Document) -> tuple[int, int]:
    fig = 0
    tab = 0
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt.startswith("图"):
            fig += 1
        if txt.startswith("表"):
            tab += 1
    return fig, tab


def main() -> None:
    src = find_source()
    backup = Path(BACKUP_NAME)
    if not backup.exists():
        shutil.copy2(src, backup)

    doc = Document(str(src))
    pattern = term_pattern(WATCH_TERMS)
    before_counts = count_hits(doc, pattern)
    before_total = sum(before_counts.values())
    before_figures, before_tables = count_figures_tables(doc)

    changed_paragraphs = replace_exact_paragraphs(doc)
    changed_cells = replace_table_cells(doc)

    out = Path(OUT_NAME)
    tmp_out = Path("decodified_output_tmp.docx")
    doc.save(tmp_out)
    shutil.copy2(tmp_out, out)

    check = Document(str(out))
    after_counts = count_hits(check, pattern)
    residuals = residual_locations(check)
    after_figures, after_tables = count_figures_tables(check)

    report = [
        "去代码化表达修改报告",
        "",
        f"源文件：{src.resolve()}",
        f"备份文件：{backup.resolve()}",
        f"输出文件：{out.resolve()}",
        "",
        f"共发现代码级表名或字段名等表达：{before_total}处",
        f"共修改正文段落：{changed_paragraphs}个",
        f"共修改表格单元格：{changed_cells}个",
        "",
        "指定关键词残留检查：",
    ]
    if residuals:
        report.append("仍有残留关键词，位置与原因如下：")
        for item in residuals:
            report.append(f"- {item}。保留原因：该处属于仍需人工复核的残留表达。")
    else:
        report.append("未发现 dim_、fact_、ads_、sys_、school_id、major_code、password_hash、app.py、dataService.js、run_full_pipeline_check.py、HTTP 200、code=0 等指定残留关键词。")

    report.extend(
        [
            "",
            "质量检查：",
            f"- 生成后的 docx 已重新打开读取：通过。",
            f"- 图题数量：修改前{before_figures}处，修改后{after_figures}处。",
            f"- 表题数量：修改前{before_tables}处，修改后{after_tables}处。",
            "- 章节编号和标题结构未做重排，参考文献列表未做主动改写。",
            "- 已重点检查物理表名、字段名、脚本文件名、接口路径和测试日志式表达。",
            "",
            "重点修改章节：",
            "- 第2章：维度建模、数据仓库和应用结果数据相关表述。",
            "- 第3章3.1.4：数据需求分析中的数据对象分类。",
            "- 第3章3.2：系统架构设计中的数据库逻辑层级。",
            "- 第3章3.3.2：数据库设计及核心数据对象表格。",
            "- 第4章4.1、4.2、4.3：工程目录、数据处理链路和算法实现表述。",
            "- 第4章4.4：后端接口服务设计。",
            "- 第4章4.6：安全审计与全链路验证。",
            "- 第5章：总结与展望中的分层数据和应用结果数据表述。",
            "",
            "修改原则说明：",
            "已将代码级实现细节改为论文级逻辑结构表述，保留了MySQL、Flask、React、Spark、ECharts、B/S架构、LSTM、FP-Growth、协同过滤和余弦相似度等关键技术路线，并保持系统设计完整性。",
            "",
            "修改后剩余广义命中统计（用于人工复核，不等同于指定残留）：",
        ]
    )
    if after_counts:
        for term, count in after_counts.most_common():
            report.append(f"- {term}: {count}")
    else:
        report.append("- 无")

    Path(REPORT_NAME).write_text("\n".join(report), encoding="utf-8")


if __name__ == "__main__":
    main()
